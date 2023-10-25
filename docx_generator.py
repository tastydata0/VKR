from docx import Document
import os
import datetime


def main():
    template_file_path = "template.docx"
    output_file_path = "result.docx"

    variables = {
        "{ПРЕДСТАВИТЕЛЬ_ФИО}": "Иванов Родитель Иванович",
        "{ПРЕДСТАВИТЕЛЬ_АДРЕС}": "г. Таганрог, Некрасовский пер, 44",
        "{РЕБЕНОК_ФИО_РОД_ПАДЕЖ}": "Иванова Ребенка Ивановича",
        "{ПРОГРАММА}": "Основы программирования на C/C++",
        "{ПРОГРАММА_ЧАСЫ}": "144",
        "{ПРОГРАММА_ЧАСЫ_АУД}": "72",
        "{РЕБЕНОК_ФИО}": "Иванов Ребенок Иванович",
        "{РЕБЕНОК_ДАТА_РОЖ}": "17.08.2009",
        "{РЕБЕНОК_МЕСТО_РОЖ}": "Ростовская область, г. Таганрог",
        "{РЕБЕНОК_МЕСТО_УЧЕБЫ}": "СОШ №8",
        "{РЕБЕНОК_КЛАСС}": "9",
        "{РЕБЕНОК_ТЕЛЕФОН}": "+7(999)987-34-43",
        "{ПРЕДСТАВИТЕЛЬ_ТЕЛЕФОН}": "+7(123)456-31-31",
        "{ПРЕДСТАВИТЕЛЬ_ПОЧТА}": "parent@gmail.com",
        "{ГОД}": datetime.datetime.now().year,
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value
                        )

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    value = str(value)
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == "__main__":
    main()
