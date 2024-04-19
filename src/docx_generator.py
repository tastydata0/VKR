from io import BytesIO
from docx import Document as DocxDocument  # type: ignore
import datetime as dt
from src.models import *
from src.name_translation import fio_to_accusative
import src.database as database
from returns.maybe import Maybe


def format_date(date: dt.date | dt.datetime) -> str:
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]

    day = date.day
    month = months[date.month - 1]
    year = date.year
    formatted_date = f"{day:02d} {month} {year} г."
    return formatted_date


def generate_doc(
    userData: User, template_name: str
) -> BytesIO:  # template_name: application, consent, ...
    template_file_path = f"data/docx_files/{template_name}.docx"

    program_info: Maybe[Program] = database.resolve_program_by_realization_id(
        userData.application.selectedProgram
    ).bind_optional(lambda program: Program(**program))

    variables = {
        "{ПРЕДСТАВИТЕЛЬ_ФИО}": userData.parentFullName,  # | "ФИО родителя",
        "{ПРЕДСТАВИТЕЛЬ_АДРЕС}": userData.parentAddress,  # | "Адрес родителя",
        "{РЕБЕНОК_ФИО_РОД_ПАДЕЖ}": fio_to_accusative(userData.fullName).value_or(
            "_" * 20
        ),  # | "ФИО ребенка",
        "{ПРОГРАММА}": program_info.bind_optional(
            lambda program: program.relevant_confirmed().formalName
        ).value_or("_" * 20),
        "{ПРОГРАММА_ЧАСЫ}": str(
            program_info.bind_optional(
                lambda program: program.relevant_confirmed().hoursAud
                + int(program.relevant_confirmed().hoursHome)
            ).value_or("__")
        ),
        "{ПРОГРАММА_ЧАСЫ_АУД}": program_info.bind_optional(
            lambda program: program.relevant_confirmed().hoursAud
        ).value_or("__"),
        "{ПРОГРАММА_НАЧАЛО}": program_info.bind_optional(
            lambda program: format_date(program.relevant_realization().realizationDate)  # type: ignore
        ).value_or("______"),
        "{ПРОГРАММА_КОНЕЦ}": program_info.bind_optional(
            lambda program: format_date(program.relevant_realization().finishDate)  # type: ignore
        ).value_or("______"),
        "{РЕБЕНОК_ФИО}": userData.fullName,  # | "ФИО ребенка",
        "{РЕБЕНОК_ДАТА_РОЖ}": userData.birthDate,  # | "Дата рождения ребенка",
        "{РЕБЕНОК_МЕСТО_РОЖ}": userData.birthPlace,  # | "Место рождения ребенка",
        "{РЕБЕНОК_МЕСТО_УЧЕБЫ}": userData.school,  # | "Школа ребенка",
        "{РЕБЕНОК_КЛАСС}": userData.schoolClass,  # | "Класс",
        "{РЕБЕНОК_ТЕЛЕФОН}": userData.phone,  # | "Телефон ребенка",
        "{ПРЕДСТАВИТЕЛЬ_ТЕЛЕФОН}": userData.parentPhone,  # | "Телефон родителя",
        "{ПРЕДСТАВИТЕЛЬ_ПОЧТА}": userData.parentEmail,  # | "Email родителя",
        "{ГОД}": datetime.now().year,
    }

    template_document = DocxDocument(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value
                        )

    docx_bytes = BytesIO()

    template_document.save(docx_bytes)

    return docx_bytes


def replace_text_in_paragraph(paragraph, key, value):
    value = str(value)
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)
